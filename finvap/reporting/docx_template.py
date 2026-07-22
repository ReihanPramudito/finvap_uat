"""Fill a custom Word (`.docx`) report template (Objective 4 / S3).

The operator authors a `.docx` whose look *and* structure are theirs; FinVAP fills
it from the assessment. Three fill mechanisms, by slot:

  * **Document Properties** — scalar identity (client/company/author/…); set once,
    cascade everywhere via Word fields. Collected by :mod:`finvap.engagement`.
  * **``{{tokens}}``** — deterministic data: severity counts, the two-tier SLA, and
    the per-finding / per-row values. Repeatable rows (scope, technical &
    compliance summaries) and the whole finding block are *cloned* per item.
  * **``[bracket]`` instructions in the Executive Summary** — the only LLM-written
    text; everything else is rendered verbatim so scores/clauses can't be
    hallucinated. Identifiers are masked before the call and the call is audited
    (S1), exactly like the Markdown report.

Findings are **grouped by vulnerability name**: one detailed section + one summary
row per unique vulnerability, listing every affected host with its own
regulatory-adjusted (fw_adj) severity; the headline shows the worst case. The
CVSS version chosen on the Setup page is displayed (and stamped into the
``CVSS_Version`` document property). Output is the filled DOCX (the editable
artifact) plus a PDF via LibreOffice headless.
"""
from __future__ import annotations

import copy
import json
import re
import subprocess
import tempfile
import time
from collections import Counter
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from lxml import etree

from .. import audit
from .deadlines import clause_for, merged_sla
from .masking import Masker
from .providers import LLMError, get_provider

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _q(tag: str) -> str:
    return f"{{{W}}}{tag}"


class TemplateError(RuntimeError):
    """Filling the template failed (missing template, bad structure, soffice absent)."""


# Shipped default template, used when neither --template nor `finvap config` names one.
DEFAULT_TEMPLATE = "VA Template.docx"


def resolve_template(name: str | None = None) -> Path:
    """Locate the template to fill.

    `name` (an explicit `--template`) wins; otherwise fall back to the `finvap
    config` template, then the bundled default. Accepts a path or a name in
    templates/ (with or without the .docx suffix).
    """
    from .. import config, settings
    if not name:
        name = (settings.load().get("template") or "").strip() or DEFAULT_TEMPLATE
    p = Path(name)
    if p.exists():
        return p
    for cand in (config.TEMPLATES_DIR / name, config.TEMPLATES_DIR / f"{name}.docx"):
        if cand.exists():
            return cand
    raise TemplateError(
        f"report template '{name}' not found (looked in {config.TEMPLATES_DIR}). "
        f"Add it to templates/, pass --template, or set one with `finvap config`.")


# Severity → (cell fill, text colour). Matches the legend the operator put in §2.3.
PALETTE = {
    "Critical": ("C00000", "FFFFFF"), "High": ("FF0000", "FFFFFF"),
    "Medium": ("FFC000", "000000"), "Low": ("92D050", "000000"),
    "Info": ("00B0F0", "000000"),
}
_SEV_ORDER = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Info": 4}
_FW_NAME = {"rmit": "Bank Negara Malaysia RMiT", "trm": "Monetary Authority of Singapore TRM"}


# --------------------------------------------------------------------------- #
# Model — grouped by vulnerability
# --------------------------------------------------------------------------- #

@dataclass
class Instance:
    ip: str
    proto: str
    port: str
    severity: str
    cvss: str
    exposure: str


@dataclass
class VulnGroup:
    name: str
    description: str = ""
    solution: str = ""
    references: list[str] = field(default_factory=list)
    cves: list[str] = field(default_factory=list)
    clauses: list[dict] = field(default_factory=list)
    instances: list[Instance] = field(default_factory=list)
    # headline (worst-case) — set after all instances are in
    severity: str = "Info"
    cvss: str = "0.0"
    vector: str = "-"
    status: str = "OPEN"
    _best: float = -1.0
    # Manual report inputs (S5.5b) — merged across the group's findings.
    poc_screenshot: str | None = None
    steps: str | None = None
    client_comments: str | None = None
    postverif_screenshot: str | None = None
    postverif_comments: str | None = None


def _disp_sev(fs) -> str:
    if fs is None:
        return "Info"
    return fs.fw_adj_severity or fs.adj_severity or fs.base_severity or "Info"


def _disp_score(fs) -> float | None:
    if fs is None:
        return None
    for v in (fs.fw_adj_score, fs.adj_score, fs.base_score):
        if v is not None:
            return v
    return None


def _split_multi(value: str | None) -> list[str]:
    if not value:
        return []
    parts = re.split(r"[\s,;]+", value.strip())
    seen, out = set(), []
    for p in parts:
        p = p.strip()
        if p and p not in seen:
            seen.add(p)
            out.append(p)
    return out


def load_model(framework: str, cvss_version: str = "4.0") -> tuple[list[VulnGroup], Counter, list, str]:
    """Build the grouped report model from the DB. Returns (groups, counts, assets, fw).

    Scores, vectors and severities are read from the requested CVSS version's layer
    (both 3.1 and 4.0 are always computed and carry their own fw_adj), so the report
    matches whichever version the operator picked on the Setup page.
    """
    from sqlmodel import select

    from ..db import get_session
    from ..models import Asset, Finding, FindingReportInput, FindingScore

    with get_session() as s:
        assets = {a.id: a for a in s.exec(select(Asset)).all()}
        findings = s.exec(select(Finding)).all()
        fs_scores = {f.finding_id: f for f in s.exec(
            select(FindingScore).where(FindingScore.cvss_version == cvss_version)).all()}
        inputs = {ri.finding_id: ri for ri in s.exec(select(FindingReportInput)).all()}

    groups: dict[str, VulnGroup] = {}
    for f in findings:
        a = assets.get(f.asset_id)
        if a is None:
            continue
        fs = fs_scores.get(f.id)
        sev = _disp_sev(fs)
        score = _disp_score(fs)
        g = groups.get(f.name)
        if g is None:
            g = groups[f.name] = VulnGroup(name=f.name)
        if not g.description:
            g.description = (f.description or f.summary or "").strip()
        if not g.solution:
            g.solution = (f.solution or "").strip()
        for ref in _split_multi(f.references):
            if ref not in g.references:
                g.references.append(ref)
        for cve in _split_multi(f.cve):
            if cve not in g.cves:
                g.cves.append(cve)
        if not g.clauses and f.regulatory_clauses:
            try:
                g.clauses = json.loads(f.regulatory_clauses)
            except ValueError:
                g.clauses = []
        # Manual report inputs (S5.5b): merge across the group's findings — first
        # non-empty value per field wins (a screenshot/note added to any instance).
        ri = inputs.get(f.id)
        if ri is not None:
            g.poc_screenshot = g.poc_screenshot or ri.poc_screenshot
            g.steps = g.steps or ri.steps
            g.client_comments = g.client_comments or ri.client_comments
            g.postverif_screenshot = g.postverif_screenshot or ri.postverif_screenshot
            g.postverif_comments = g.postverif_comments or ri.postverif_comments
        g.instances.append(Instance(
            ip=a.ip_address, proto=(f.protocol or "tcp").upper(),
            port=str(f.port) if f.port else "", severity=sev,
            cvss=f"{score:.1f}" if score is not None else "0.0",
            exposure=a.exposure.value,
        ))
        # headline = worst (highest-scoring) instance
        s_val = score if score is not None else -1.0
        rank = (-_SEV_ORDER.get(sev, 4), s_val)
        if rank > (-_SEV_ORDER.get(g.severity, 4), g._best):
            g.severity = sev
            g.cvss = f"{score:.1f}" if score is not None else "0.0"
            g.vector = (fs.adj_vector if fs and fs.adj_vector else "-")
            g._best = s_val

    ordered = sorted(groups.values(),
                     key=lambda g: (_SEV_ORDER.get(g.severity, 4), -g._best, g.name.lower()))
    counts: Counter = Counter()
    for g in ordered:
        counts[g.severity] += 1
    return ordered, counts, list(assets.values()), framework


# --------------------------------------------------------------------------- #
# OXML helpers
# --------------------------------------------------------------------------- #

def _para_text(p) -> str:
    return "".join((t.text or "") for t in p.iter(_q("t")))


def _unhighlight_para(p) -> None:
    """Remove the yellow authoring highlight from every run in a paragraph — filled
    content should look clean, not marked-up (the highlight only flags fill points)."""
    for r in p.findall(_q("r")):
        rpr = r.find(_q("rPr"))
        if rpr is not None:
            for h in rpr.findall(_q("highlight")):
                rpr.remove(h)


def _title_env(value: str) -> str:
    """Display an environment tag title-cased (production -> Production, uat -> UAT)."""
    return "UAT" if (value or "").lower() == "uat" else (value or "").capitalize()


def _replace_in_para(p, mapping: dict) -> None:
    """Replace ``{{tokens}}`` in a paragraph, even when split across runs.

    Substitutes within each contiguous run of *plain* text, treating field regions
    (``fldSimple`` and complex ``fldChar`` fields, e.g. inline DOCPROPERTY) as
    boundaries. A field's runs are left untouched, so a resolved property value is
    never flattened into the sentence *and* left behind for the consumer to
    re-render — which previously duplicated the client/company short-name at the
    end of the paragraph. Untouched if no key is present.
    """
    segments: list[list] = []      # each = the <w:t> nodes of one plain-text run of the paragraph
    current: list = []
    field_depth = 0                # >0 while inside a complex fldChar begin…end field

    def _flush():
        nonlocal current
        if current:
            segments.append(current)
            current = []

    for child in list(p):
        if child.tag == _q("fldSimple"):
            _flush()               # a simple field — its value is resolved in place, leave it
            continue
        if child.tag != _q("r"):
            continue               # pPr, bookmarks, etc.
        fld = child.find(_q("fldChar"))
        if fld is not None:        # complex-field boundary marker
            ftype = fld.get(_q("fldCharType"))
            if ftype == "begin":
                _flush(); field_depth += 1
            elif ftype == "end":
                field_depth = max(0, field_depth - 1); _flush()
            continue
        if child.find(_q("instrText")) is not None or field_depth > 0:
            continue               # a field's instruction or result run — leave it alone
        current.extend(child.findall(_q("t")))   # a plain run: collect its text nodes
    _flush()

    changed = False
    for seg in segments:
        full = "".join(t.text or "" for t in seg)
        if not any(k in full for k in mapping):
            continue
        for k, v in mapping.items():
            full = full.replace(k, v)
        seg[0].text = full
        for t in seg[1:]:
            t.text = ""
        changed = True
    if changed:
        _unhighlight_para(p)


def _replace_in_elements(elements, mapping: dict) -> None:
    for el in elements:
        for p in el.iter(_q("p")):
            _replace_in_para(p, mapping)


def _shade_cell(tc, fill: str, text: str) -> None:
    """Apply a background fill to a table cell and set its runs' font colour."""
    tcPr = tc.find(_q("tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, _q("tcPr"))
        tc.insert(0, tcPr)
    for old in tcPr.findall(_q("shd")):
        tcPr.remove(old)
    shd = etree.SubElement(tcPr, _q("shd"))
    shd.set(_q("val"), "clear")
    shd.set(_q("color"), "auto")
    shd.set(_q("fill"), fill)
    for r in tc.iter(_q("r")):
        rPr = r.find(_q("rPr"))
        if rPr is None:
            rPr = etree.Element(_q("rPr"))
            r.insert(0, rPr)
        for old in rPr.findall(_q("color")):
            rPr.remove(old)
        col = etree.SubElement(rPr, _q("color"))
        col.set(_q("val"), text)


def _shade_sev_cell_in_row(tr, sev: str, col_idx: int) -> None:
    cells = tr.findall(_q("tc"))
    if col_idx < len(cells) and sev in PALETTE:
        fill, txt = PALETTE[sev]
        _shade_cell(cells[col_idx], fill, txt)


def _cell_text(tc) -> str:
    return "".join((t.text or "") for t in tc.iter(_q("t")))


def _find_block_element(elements, token: str):
    """First element (paragraph or table) in `elements` whose text holds `token`."""
    for el in elements:
        if token in "".join((t.text or "") for t in el.iter(_q("t"))):
            return el
    return None


def _repeat_paragraph(template_p, fills: list[dict]) -> None:
    """Clone a token paragraph once per fill dict (in place); blank if none."""
    if not fills:
        _replace_in_para(template_p, {k: "—" for k in _ONLY_BRACKETS_OR_TOKENS(template_p)})
        return
    for fl in fills:
        clone = copy.deepcopy(template_p)
        _replace_in_para(clone, fl)
        template_p.addprevious(clone)
    template_p.getparent().remove(template_p)


def _ONLY_BRACKETS_OR_TOKENS(p) -> list[str]:
    return re.findall(r"\{\{[^}]+\}\}", _para_text(p))


def _repeat_table_row(tbl, fills: list[dict], *, sev_key: str | None = None,
                      sev_col: int | None = None) -> None:
    """Clone the template row (row index 1) once per fill; shade by severity if asked."""
    rows = tbl.findall(_q("tr"))
    if len(rows) < 2:
        return
    tmpl = rows[1]
    for fl in fills:
        clone = copy.deepcopy(tmpl)
        for p in clone.iter(_q("p")):
            _replace_in_para(p, fl)
        if sev_key and sev_col is not None:
            _shade_sev_cell_in_row(clone, fl.get(sev_key, ""), sev_col)
        tmpl.addprevious(clone)
    tmpl.getparent().remove(tmpl)


# --------------------------------------------------------------------------- #
# Document Properties
# --------------------------------------------------------------------------- #

_DOCPROP_RE = re.compile(r'DOCPROPERTY\s+"?([A-Za-z0-9_]+)"?', re.I)


def _update_custom_xml(doc, values: dict) -> None:
    """Set the stored custom-property values (the source Word fields read on refresh)."""
    for part in doc.part.package.iter_parts():
        if str(part.partname) != "/docProps/custom.xml":
            continue
        root = etree.fromstring(part.blob)
        ns = root.nsmap.get(None, "")
        vt = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
        for prop in root.findall(f"{{{ns}}}property"):
            name = prop.get("name")
            if name in values:
                lp = prop.find(f"{{{vt}}}lpwstr")
                if lp is not None:
                    lp.text = values[name]
        part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8",
                                    standalone=True)
        return


def _refresh_docprop_cache(p, values: dict) -> None:
    """Overwrite the cached result text of DOCPROPERTY fields so the displayed value
    is correct even if the consumer never refreshes fields."""
    # Simple fields.
    for fs in p.findall(".//" + _q("fldSimple")):
        m = _DOCPROP_RE.search(fs.get(_q("instr")) or "")
        if m and m.group(1) in values:
            ts = list(fs.iter(_q("t")))
            if ts:
                ts[0].text = values[m.group(1)]
                for t in ts[1:]:
                    t.text = ""
    # Complex fields (begin / instrText / separate / result / end).
    state = None
    prop = None
    first = True
    instr = ""
    for r in p.findall(_q("r")):
        fld = r.find(_q("fldChar"))
        if fld is not None:
            t = fld.get(_q("fldCharType"))
            if t == "begin":
                state, prop, first, instr = "instr", None, True, ""
            elif t == "separate":
                m = _DOCPROP_RE.search(instr)
                prop = m.group(1) if m else None
                state = "result"
            elif t == "end":
                state, prop = None, None
            continue
        it = r.find(_q("instrText"))
        if it is not None and state == "instr":
            instr += it.text or ""
            continue
        if state == "result" and prop in values:
            ts = r.findall(_q("t"))
            if first and ts:
                ts[0].text = values[prop]
                for t in ts[1:]:
                    t.text = ""
                first = False
            else:
                for t in ts:
                    t.text = ""


def _set_docproperties(doc, values: dict) -> None:
    values = {k: (v if v is not None else "") for k, v in values.items()}
    _update_custom_xml(doc, values)
    for part_el in _all_story_roots(doc):
        for p in part_el.iter(_q("p")):
            _refresh_docprop_cache(p, values)


def _all_story_roots(doc):
    """Body + every header/footer element root (DOCPROPERTY fields can live in any)."""
    roots = [doc.element.body]
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            try:
                roots.append(hf._element)
            except Exception:
                pass
    return roots


def _set_update_fields(doc) -> None:
    """Ask the consumer (Word / LibreOffice) to refresh fields (TOC, page #) on load."""
    settings = doc.settings.element
    if settings.find(_q("updateFields")) is None:
        uf = etree.SubElement(settings, _q("updateFields"))
        uf.set(_q("val"), "true")


# --------------------------------------------------------------------------- #
# Executive-summary prose (the only LLM text)
# --------------------------------------------------------------------------- #

_SYS_EXEC = (
    "You are a senior financial-sector security auditor completing the executive "
    "summary of a vulnerability assessment for non-technical senior management. Use "
    "ONLY the facts provided — never invent vulnerabilities, regulations, scores or "
    "CVEs. Each value must be a sentence FRAGMENT that grammatically continues the "
    "lead-in shown, with NO leading capital (unless a proper noun/acronym) and NO "
    "trailing full stop. Reply with STRICT JSON: {"
    "\"posture\": <one of 'possesses a strong security posture' / 'requires moderate "
    "security improvement' / 'is highly vulnerable to compromise'>, "
    "\"core_issues\": <the CONTINUATION only after 'the risks identified involve', e.g. "
    "'outdated software components and weak authentication controls' — do NOT repeat the "
    "words 'risks identified involve'>, "
    "\"business_impacts\": <the CONTINUATION only after 'allow an attacker to', e.g. "
    "'access sensitive data or move laterally into the network' — do NOT repeat the words "
    "'allow an attacker to'>, "
    "\"recommendations\": [<3-5 short action-oriented goals, highest severity first>]}. "
    "No markdown, no extra keys."
)

# The model sometimes echoes the lead-in; strip it so the sentence reads once.
_CORE_LEADIN = re.compile(r"^(the\s+)?(most\s+significant\s+)?risks?\s+(identified\s+)?"
                          r"(that\s+)?involve\s+", re.I)
_IMPACT_LEADIN = re.compile(r"^(that\s+|which\s+)?(would\s+|could\s+|can\s+)?(allow\s+)?"
                            r"(an?\s+)?attacker\s+(to|could|can|would)\s+(to\s+)?", re.I)


def _frag(s: str, leadin: re.Pattern | None = None) -> str:
    """Coerce an LLM value into a mid-sentence fragment: drop any echoed lead-in, a
    trailing period, and lowercase the first char unless it's a proper noun/acronym."""
    s = (s or "").strip()
    if leadin is not None:
        s = leadin.sub("", s).strip()
    s = s.rstrip(".").strip()
    if not s:
        return s
    first = s.split(maxsplit=1)[0]
    if not (len(first) > 1 and first.isupper()):  # keep 'FTP…', 'SSH…', 'TLS…'
        s = s[0].lower() + s[1:]
    return s


def _exec_facts(masker: Masker, groups, counts, fw, target) -> str:
    fwname = _FW_NAME.get(fw, fw)
    lines = [f"Network assessed: {masker.mask(target)}", f"Regulatory framework: {fwname}",
             "Unique vulnerabilities by severity (regulatory-adjusted): " +
             ", ".join(f"{k} {counts.get(k, 0)}" for k in ("Critical", "High", "Medium", "Low", "Info")),
             "", "Most severe vulnerabilities (with the regulatory clause they implicate):"]
    for g in groups[:12]:
        cite = g.clauses[0]["citation"] if g.clauses else "no mapped clause"
        hosts = len(g.instances)
        lines.append(f"- [{g.severity}] {masker.mask(g.name)} on {hosts} host(s) -> {cite}")
    return "\n".join(lines)


def _deterministic_prose(counts) -> dict:
    nonzero = [(k, counts.get(k, 0)) for k in ("Critical", "High", "Medium", "Low", "Info")
               if counts.get(k, 0)]
    top = nonzero[0][0] if nonzero else "Info"
    posture = ("is highly vulnerable to compromise" if top in ("Critical", "High")
               else "requires moderate security improvement" if top == "Medium"
               else "possesses a strong security posture")
    return {
        "posture": posture,
        "core_issues": "outdated or misconfigured software components and weak security controls",
        "business_impacts": "compromise sensitive data, gain unauthorised access, or disrupt services",
        "recommendations": [
            "Remediate the highest-severity findings first, following the timelines in this report.",
            "Patch outdated and end-of-life software components across the affected assets.",
            "Harden weak configurations and enforce strong authentication and encryption.",
        ],
    }


def _exec_prose(groups, counts, fw, target, provider, model) -> dict:
    masker = Masker()
    facts = _exec_facts(masker, groups, counts, fw, target)
    prov = get_provider(provider, model)
    ok, _ = prov.available()
    if not ok or prov.name == "template":
        audit.record("report.prose", target="exec_template", status="ok",
                     summary="executive summary: deterministic prose (no LLM)")
        return _deterministic_prose(counts)
    t0 = time.time()
    try:
        out = prov.complete(_SYS_EXEC, facts)
    except LLMError as e:
        audit.record("llm.call", target="exec_summary", status="error",
                     summary=f"{prov.name} exec_summary failed: {e}",
                     duration_ms=int((time.time() - t0) * 1000))
        return _deterministic_prose(counts)
    audit.ai_call(stage="exec_summary", provider=prov.name, model=getattr(prov, "model", ""),
                  system=_SYS_EXEC, user_sent=facts, response_raw=out or "",
                  placeholder_map=masker.map, response_unmasked=masker.unmask(out or ""),
                  duration_ms=int((time.time() - t0) * 1000))
    try:
        m = re.search(r"\{.*\}", out or "", re.S)
        data = json.loads(m.group(0)) if m else {}
        det = _deterministic_prose(counts)
        return {
            "posture": str(data.get("posture") or det["posture"]),
            "core_issues": str(data.get("core_issues") or det["core_issues"]),
            "business_impacts": str(data.get("business_impacts") or det["business_impacts"]),
            "recommendations": [str(x) for x in data.get("recommendations") or []]
                               or det["recommendations"],
        }
    except (ValueError, TypeError):
        return _deterministic_prose(counts)


def _fill_exec_summary(doc, prose: dict) -> None:
    body = doc.element.body
    # Narrative paragraph: the three [brackets] in order = posture, core issues, impact.
    narrative = None
    rec_bullet = None
    for p in body.iter(_q("p")):
        txt = _para_text(p)
        if "[Choose one" in txt:
            narrative = p
        if "[Action-oriented remediation goal" in txt:
            rec_bullet = p
    if narrative is not None:
        # posture already fits "the target <posture>"; the other two must be fragments
        # that continue the lead-in without echoing it.
        order = [prose["posture"],
                 _frag(prose["core_issues"], _CORE_LEADIN),
                 _frag(prose["business_impacts"], _IMPACT_LEADIN)]
        full = _para_text(narrative)
        for val in order:
            full = re.sub(r"\[[^\]]*\]", lambda _m, v=val: v, full, count=1)
        tnodes = list(narrative.iter(_q("t")))
        if tnodes:
            tnodes[0].text = full
            for t in tnodes[1:]:
                t.text = ""
            _unhighlight_para(narrative)
    if rec_bullet is not None:
        goals = prose["recommendations"] or ["Remediate the highest-severity findings first."]
        for goal in goals:
            clone = copy.deepcopy(rec_bullet)
            tnodes = list(clone.iter(_q("t")))
            if tnodes:
                tnodes[0].text = goal
                for t in tnodes[1:]:
                    t.text = ""
                _unhighlight_para(clone)
            rec_bullet.addprevious(clone)
        rec_bullet.getparent().remove(rec_bullet)


# --------------------------------------------------------------------------- #
# Per-finding prose (LLM rewrite of the scanner's description + remediation)
# --------------------------------------------------------------------------- #

_SYS_FINDING = (
    "You are a security auditor editing a vulnerability report. Rewrite the DESCRIPTION "
    "and REMEDIATION below into clear, concise, professional prose for a client report. "
    "Use ONLY the information given — do NOT invent CVEs, version numbers, hostnames, "
    "ports, IPs, or any fact not present, and preserve any specific values that ARE "
    "present. Reply with STRICT JSON {\"description\": \"…\", \"recommendation\": \"…\"}. "
    "No markdown, no extra keys."
)


# --------------------------------------------------------------------------- #
# Finding block (the repeatable detailed section)
# --------------------------------------------------------------------------- #

def _finding_block_elements(doc):
    """Body elements making up the one template finding block: from the Heading 2
    paragraph carrying {{F_TITLE}} to the last element before the body's sectPr."""
    body = doc.element.body
    kids = list(body.iterchildren())
    start = None
    for i, el in enumerate(kids):
        if el.tag != _q("p"):
            continue
        ppr = el.find(_q("pPr"))
        style = ppr.find(_q("pStyle")) if ppr is not None else None
        is_h2 = style is not None and style.get(_q("val")) == "Heading2"
        if is_h2 and "{{F_TITLE}}" in _para_text(el):
            start = i
            break
    if start is None:
        raise TemplateError("template has no Heading-2 '{{F_TITLE}}' finding block")
    end = len(kids)
    for j in range(len(kids) - 1, start, -1):
        if kids[j].tag == _q("sectPr"):
            end = j
            break
    return kids[start:end]


_NA_TEXT = "N/A — non-intrusive assessment; no exploitation performed."


def _fill_screenshot(doc, block_elements, token: str, path) -> None:
    """Embed a screenshot at the token's paragraph, or replace the token with an
    N/A note when there's no (valid) image. Validates the image *before* touching
    the token, so a missing/corrupt file degrades cleanly to the note."""
    p_el = _find_block_element(block_elements, token)
    if p_el is None:
        return
    valid = False
    if path and Path(path).exists():
        try:
            from docx.image.image import Image
            Image.from_file(str(path))          # parse-check; raises on a bad image
            valid = True
        except Exception:                       # pragma: no cover - corrupt image
            valid = False
    if not valid:
        _replace_in_elements([p_el], {token: _NA_TEXT})
        return
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph
    for t in p_el.iter(_q("t")):                # clear the token text, keep the paragraph
        t.text = ""
    _unhighlight_para(p_el)
    try:
        Paragraph(p_el, doc).add_run().add_picture(str(path), width=Inches(5.8))
    except Exception:                           # pragma: no cover - shouldn't happen post-validate
        Paragraph(p_el, doc).add_run(f"[screenshot unavailable: {Path(path).name}]")


def _fill_one_finding(doc, block_elements, g: VulnGroup, framework: str, ref: str) -> None:
    """Fill a deep-copied finding block (list of elements) for one vuln group."""
    # 1) repeats first (so scalar replace below also covers the clones)
    asset_p = _find_block_element(block_elements, "{{F_ASSET_IP}}")
    if asset_p is not None:
        fills = [{
            "{{F_ASSET_IP}}": ins.ip, "{{F_ASSET_PROTO}}": ins.proto,
            "{{F_ASSET_PORT}}": ins.port or "—", "{{F_ASSET_CVSS}}": ins.cvss,
            "{{F_ASSET_SEVERITY}}": ins.severity,
        } for ins in g.instances]
        _repeat_paragraph(asset_p, fills)

    cve_p = _find_block_element(block_elements, "{{F_CVE}}")
    if cve_p is not None:
        _repeat_paragraph(cve_p, [{"{{F_CVE}}": c} for c in g.cves] or [{"{{F_CVE}}": "None"}])

    ref_p = _find_block_element(block_elements, "{{F_REF_LINK}}")
    if ref_p is not None:
        _repeat_paragraph(ref_p, [{"{{F_REF_LINK}}": r} for r in g.references]
                          or [{"{{F_REF_LINK}}": "None provided"}])

    # severity table (single row) + shading
    for el in block_elements:
        if el.tag == _q("tbl") and "{{F_SEVERITY}}" in "".join(
                (t.text or "") for t in el.iter(_q("t"))):
            rows = el.findall(_q("tr"))
            if len(rows) >= 2:
                for p in rows[1].iter(_q("p")):
                    _replace_in_para(p, {"{{F_SEVERITY}}": g.severity,
                                         "{{F_CVSS}}": g.cvss, "{{F_VECTOR}}": g.vector})
                _shade_sev_cell_in_row(rows[1], g.severity, 0)

    # clause table (repeat row per cited clause)
    for el in block_elements:
        if el.tag == _q("tbl") and "{{F_CLAUSE_ID}}" in "".join(
                (t.text or "") for t in el.iter(_q("t"))):
            if g.clauses:
                fills = [{
                    "{{F_CLAUSE_ID}}": _norm_cite(c.get("citation", "")),
                    "{{F_CLAUSE_TITLE}}": c.get("section", ""),
                    "{{F_CLAUSE_RATIONALE}}": c.get("reason") or (c.get("excerpt") or "")[:200],
                } for c in g.clauses]
                _repeat_table_row(el, fills)
            else:
                # No clause: the cell literal is "{{F_CLAUSE_ID}} - {{F_CLAUSE_TITLE}}", so
                # fill blanks then set the first cell cleanly (no stray "— - ").
                _repeat_table_row(el, [{"{{F_CLAUSE_ID}}": "", "{{F_CLAUSE_TITLE}}": "",
                                        "{{F_CLAUSE_RATIONALE}}": "General security best practice."}])
                row = el.findall(_q("tr"))[1]
                cell0 = row.findall(_q("tc"))[0]
                ps = cell0.findall(_q("p"))
                if ps:
                    tnodes = list(ps[0].iter(_q("t")))
                    if tnodes:
                        tnodes[0].text = "No specific clause mapped"
                        for t in tnodes[1:]:
                            t.text = ""
                        _unhighlight_para(ps[0])

    # 2) scalars across the whole block
    _replace_in_elements(block_elements, {
        "{{F_TITLE}}": g.name,
        "{{F_DESCRIPTION}}": g.description or "—",
        "{{F_RECOMMENDATION}}": g.solution or "—",
        "{{F_STATUS}}": g.status,
        "{{F_CLIENT_COMMENTS}}": g.client_comments or "—",
        "{{F_POSTVERIF_COMMENTS}}": g.postverif_comments or "—",
    })

    # 3) screenshots (image embed, or an N/A note when none)
    _fill_screenshot(doc, block_elements, "{{F_POC_SCREENSHOT}}", g.poc_screenshot)
    _fill_screenshot(doc, block_elements, "{{F_POSTVERIF_SCREENSHOT}}", g.postverif_screenshot)


# --------------------------------------------------------------------------- #
# Numbered-list restart (so each finding's "Steps to Reproduce" starts at 1)
# --------------------------------------------------------------------------- #

def _numbering_root(doc):
    try:
        return doc.part.numbering_part.element
    except Exception:
        return None


def _ordered_num_ids(doc) -> set:
    """numIds whose format is ordered (decimal/letter/roman) — these show a number and
    must restart per finding; bullets don't."""
    root = _numbering_root(doc)
    if root is None:
        return set()
    ordered_abs = set()
    for an in root.findall(_q("abstractNum")):
        lvl = an.find(_q("lvl"))
        fmt = lvl.find(_q("numFmt")) if lvl is not None else None
        if fmt is not None and (fmt.get(_q("val")) or "") not in ("bullet", "none", ""):
            ordered_abs.add(an.get(_q("abstractNumId")))
    return {num.get(_q("numId")) for num in root.findall(_q("num"))
            if (num.find(_q("abstractNumId")) is not None
                and num.find(_q("abstractNumId")).get(_q("val")) in ordered_abs)}


def _new_num(root, base_numid: str):
    """Add a fresh <w:num> reusing base_numid's abstractNum but with an explicit
    level-0 startOverride=1 — an independent instance that restarts at 1 *and* survives
    a LibreOffice re-save (identical no-override duplicates get merged back). Returns
    the new numId (str) or None."""
    aid = None
    for num in root.findall(_q("num")):
        if num.get(_q("numId")) == base_numid:
            an = num.find(_q("abstractNumId"))
            aid = an.get(_q("val")) if an is not None else None
            break
    if aid is None:
        return None
    used = [int(n.get(_q("numId"))) for n in root.findall(_q("num"))
            if (n.get(_q("numId")) or "").isdigit()]
    newid = str((max(used) if used else 0) + 1)
    num = etree.SubElement(root, _q("num"))
    num.set(_q("numId"), newid)
    etree.SubElement(num, _q("abstractNumId")).set(_q("val"), aid)
    lo = etree.SubElement(num, _q("lvlOverride"))
    lo.set(_q("ilvl"), "0")
    etree.SubElement(lo, _q("startOverride")).set(_q("val"), "1")
    return newid


def _restart_ordered_lists(doc, block_elements, ordered_nums) -> None:
    """Give this finding block its own instance of each ordered list so numbering
    (e.g. Steps to Reproduce) restarts at 1 instead of continuing across findings."""
    root = _numbering_root(doc)
    if root is None or not ordered_nums:
        return
    remap: dict = {}
    for el in block_elements:
        for p in el.iter(_q("p")):
            ppr = p.find(_q("pPr"))
            numpr = ppr.find(_q("numPr")) if ppr is not None else None
            numid = numpr.find(_q("numId")) if numpr is not None else None
            if numid is None or numid.get(_q("val")) not in ordered_nums:
                continue
            base = numid.get(_q("val"))
            remap.setdefault(base, _new_num(root, base) or base)
            numid.set(_q("val"), remap[base])


# --------------------------------------------------------------------------- #
# Top-level
# --------------------------------------------------------------------------- #

def _global_tokens(counts, fw, sla) -> dict:
    total = sum(counts.values())
    nonzero = [(k, counts.get(k, 0)) for k in ("Critical", "High", "Medium", "Low", "Info")
               if counts.get(k, 0)]
    breakdown = ", ".join(f"{n} {k}" for k, n in nonzero[:-1])
    if len(nonzero) > 1:
        breakdown += f" and {nonzero[-1][1]} {nonzero[-1][0]}"
    elif nonzero:
        breakdown = f"{nonzero[0][1]} {nonzero[0][0]}"
    else:
        breakdown = "no"
    tok = {
        "{{TOTAL_FINDINGS}}": str(total),
        "{{N_CRITICAL}}": str(counts.get("Critical", 0)),
        "{{N_HIGH}}": str(counts.get("High", 0)),
        "{{N_MEDIUM}}": str(counts.get("Medium", 0)),
        "{{N_LOW}}": str(counts.get("Low", 0)),
        "{{N_INFO}}": str(counts.get("Info", 0)),
        "{{SEVERITY_BREAKDOWN}}": f"{breakdown} finding(s)",
        "{{DEADLINE_MANDATE_CLAUSE}}": clause_for(fw) or "The applicable regulation",
    }
    for sev, key in (("Critical", "CRITICAL"), ("High", "HIGH"),
                     ("Medium", "MEDIUM"), ("Low", "LOW")):
        tok[f"{{{{SLA_{key}_EXT}}}}"] = str(sla[sev]["ext"])
        tok[f"{{{{SLA_{key}_INT}}}}"] = str(sla[sev]["int"])
    return tok


def _norm_cite(text: str) -> str:
    """Display-normalise a clause citation (RMIT -> RMiT). Lazy import avoids a
    compliance<->reporting import cycle."""
    from ..compliance import normalize_citation
    return normalize_citation(text)


def fill(template_path, out_base, *, framework: str = "rmit", cvss_version: str = "4.0",
         metadata: dict | None = None, target: str | None = None, provider: str | None = None,
         model: str | None = None, sla_overrides: dict | None = None, pdf: bool = True,
         progress=None) -> list[Path]:
    """Fill `template_path` from the DB and write a DOCX (+ PDF). Returns the paths.

    `progress` is an optional callback(message) for the CLI spinner — the per-finding
    LLM rewrites can take a while on a local model.
    """
    def _say(msg: str) -> None:
        if progress:
            progress(msg)

    try:
        from docx import Document
    except ImportError as e:  # pragma: no cover
        raise TemplateError(f"python-docx not installed ({e}); `pip install -e .`") from e

    template_path = Path(template_path)
    if not template_path.exists():
        raise TemplateError(f"template not found: {template_path}")
    doc = Document(str(template_path))

    groups, counts, assets, fw = load_model(framework, cvss_version)
    sla = merged_sla(sla_overrides)

    # 1) Document Properties (identity + framework name + CVSS version).
    props = {k: v for k, v in (metadata or {}).items() if not k.startswith("__")}
    props["Framework"] = _FW_NAME.get(fw, fw)
    props["CVSS_Version"] = cvss_version
    _set_docproperties(doc, props)

    # 2) LLM prose (masked + audited): the executive summary. Per-finding
    #    description/recommendation are already persisted by the `rewrite_findings`
    #    pipeline step (so they're editable), so we render the stored text verbatim.
    _say("writing executive summary…")
    prose = _exec_prose(groups, counts, fw, target or "the assessed network", provider, model)
    _fill_exec_summary(doc, prose)

    # 3) Repeatable tables: scope (all assets), technical & compliance summaries.
    body = doc.element.body
    for tbl in body.iter(_q("tbl")):
        head = "".join((t.text or "") for t in tbl.findall(_q("tr"))[0].iter(_q("t"))) if \
            tbl.findall(_q("tr")) else ""
        sample = "".join((t.text or "") for t in tbl.iter(_q("t")))
        if "{{ASSET_IP}}" in sample:
            _repeat_table_row(tbl, [{"{{ASSET_IP}}": a.ip_address,
                                     "{{ASSET_ENV}}": _title_env(a.environment.value)} for a in assets]
                              or [{"{{ASSET_IP}}": "—", "{{ASSET_ENV}}": "—"}])
        elif "{{VULN_NAME}}" in sample and "{{STATUS}}" in sample:  # technical summary
            _repeat_table_row(tbl, [{
                "{{REF}}": f"4.{i}", "{{SEVERITY}}": g.severity, "{{CVSS}}": g.cvss,
                "{{VULN_NAME}}": g.name, "{{STATUS}}": g.status,
            } for i, g in enumerate(groups, 1)], sev_key="{{SEVERITY}}", sev_col=1)
        elif "{{CLAUSES}}" in sample:  # compliance mapping
            _repeat_table_row(tbl, [{
                "{{REF}}": f"4.{i}", "{{SEVERITY}}": g.severity, "{{VULN_NAME}}": g.name,
                "{{CLAUSES}}": _norm_cite(", ".join(c.get("citation", "") for c in g.clauses)) or "—",
            } for i, g in enumerate(groups, 1)], sev_key="{{SEVERITY}}", sev_col=1)

    # 4) Finding block: clone per vuln group. Insert each clone (attached to the
    #    body) *before* filling, so the in-block repeats can add/remove siblings.
    ordered_nums = _ordered_num_ids(doc)
    tmpl_block = _finding_block_elements(doc)
    for i, g in enumerate(groups, 1):
        inserted = []
        for el in tmpl_block:
            c = copy.deepcopy(el)
            tmpl_block[0].addprevious(c)
            inserted.append(c)
        _fill_one_finding(doc, inserted, g, fw, f"4.{i}")
        _restart_ordered_lists(doc, inserted, ordered_nums)  # Steps restart at 1 per finding
        # Reproduction steps AFTER the restart, so each cloned list item inherits
        # this finding's fresh numId (one numbered item per non-empty line).
        steps_p = _find_block_element(inserted, "{{F_STEPS}}")
        if steps_p is not None:
            lines = [ln.strip() for ln in (g.steps or "").splitlines() if ln.strip()]
            _repeat_paragraph(steps_p, [{"{{F_STEPS}}": ln} for ln in lines]
                              or [{"{{F_STEPS}}": _NA_TEXT}])
    for el in tmpl_block:
        el.getparent().remove(el)

    # 5) Global scalar tokens (counts, SLA, mandate clause) + cover fields (S5.5b).
    cover = {
        "{{DRAFT_FINAL}}": (metadata or {}).get("__draft_final") or "Draft",
        "{{VERSION}}": (metadata or {}).get("__version") or "0.1",
        "{{COVER_DATE}}": ((metadata or {}).get("__cover_date")
                           or (metadata or {}).get("Date_DraftReport") or ""),
    }
    _replace_in_elements([body], {**_global_tokens(counts, fw, sla), **cover})
    _set_update_fields(doc)

    out_base = Path(out_base)
    out_base.parent.mkdir(parents=True, exist_ok=True)
    docx_path = out_base.with_suffix(".docx")
    doc.save(str(docx_path))
    _say("finalising (refreshing the table of contents)…")
    written = [docx_path]
    extra = _finalize(docx_path, pdf=pdf)  # refreshes the DOCX in place; returns the PDF
    if extra is not None:
        written.append(extra)
    return written


def _uno_python() -> str | None:
    """An interpreter that can `import uno` (often the system python, not the venv)."""
    import shutil
    import sys
    seen = set()
    for cand in (sys.executable, "/usr/bin/python3", shutil.which("python3"),
                 "/usr/lib/libreoffice/program/python"):
        if not cand or cand in seen:
            continue
        seen.add(cand)
        try:
            subprocess.run([cand, "-c", "import uno"], check=True, capture_output=True, timeout=20)
            return cand
        except (subprocess.SubprocessError, OSError):
            continue
    return None


def _finalize(docx_path: Path, *, pdf: bool) -> Path | None:
    """Refresh the filled DOCX's fields/TOC in place and, if `pdf`, also export a PDF.

    Prefers the UNO route (one LibreOffice pass that updates the TOC index, saves the
    refreshed DOCX back, and exports the PDF). Falls back to a plain headless convert
    for the PDF (valid, but TOC stays as authored); the DOCX then keeps its
    `updateFields` flag so Word refreshes it on open. Returns the PDF path or None."""
    import shutil
    out_pdf = docx_path.with_suffix(".pdf")

    py = _uno_python()
    converter = Path(__file__).with_name("_soffice_convert.py")
    if py and converter.exists():
        try:
            subprocess.run([py, str(converter), str(docx_path), (str(out_pdf) if pdf else "-")],
                           check=True, capture_output=True, timeout=300)
            if not pdf:
                return None
            if out_pdf.exists():
                return out_pdf
        except (subprocess.SubprocessError, OSError):
            pass  # fall through to the plain convert (or just keep the DOCX)

    if not pdf:
        return None  # DOCX kept as written (updateFields set — refreshes on open)

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise TemplateError("LibreOffice (soffice) not found — needed to render the report PDF; "
                            "the DOCX was still written.")
    with tempfile.TemporaryDirectory() as profile:
        try:
            subprocess.run(
                [soffice, "--headless", f"-env:UserInstallation=file://{profile}",
                 "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                check=True, capture_output=True, timeout=180,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
            raise TemplateError(f"LibreOffice PDF conversion failed: {e}") from e
    if not out_pdf.exists():
        raise TemplateError("LibreOffice ran but produced no PDF; the DOCX was still written.")
    return out_pdf
