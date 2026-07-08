"""Tests for the custom .docx report template path (S3).

Covers the engagement-metadata store, the two-tier remediation SLA, and the
docx fill engine — grouping by vulnerability, Document-Property resolution,
severity shading, and that no fill tokens are left behind (bar the TOC field,
which a consumer refreshes).
"""
import json
import re

from docx import Document
from docx.oxml.ns import qn
from sqlmodel import Session, SQLModel, create_engine

from finvap import db, engagement
from finvap import settings as user_settings
from finvap.models import Asset, Environment, Exposure, Finding, FindingScore
from finvap.reporting import deadlines, docx_template

TEMPLATE = "templates/VA Template.docx"
_VEC = "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:H/VI:H/VA:H/SC:N/SI:N/SA:N"


# --- engagement metadata -----------------------------------------------------

def test_engagement_save_load_clear():
    engagement.save({"Client_Full_Name": "Acme Bank", "bogus": "x"})
    assert engagement.load() == {"Client_Full_Name": "Acme Bank"}  # unknown key dropped
    assert engagement.clear() is True
    assert engagement.load() == {} and engagement.clear() is False


# --- two-tier SLA ------------------------------------------------------------

def test_merged_sla_applies_overrides_only_where_valid():
    sla = deadlines.merged_sla({"Critical": {"ext": 3}, "Bogus": {"ext": 1}, "Low": {"int": -5}})
    assert sla["Critical"]["ext"] == 3
    assert sla["Critical"]["int"] == deadlines.DEFAULT_SLA["Critical"]["int"]  # untouched
    assert sla["Low"]["int"] == deadlines.DEFAULT_SLA["Low"]["int"]            # invalid ignored
    assert "Bogus" not in sla


def test_settings_sla_roundtrip_preserves_preferences():
    user_settings.save({"framework": "trm"})
    user_settings.save_sla({"High": {"ext": 21, "int": 45}})
    assert user_settings.load()["framework"] == "trm"             # prefs preserved
    assert user_settings.load_sla()["High"]["ext"] == 21
    assert user_settings.reset_sla() is True
    assert user_settings.load_sla() == {} and user_settings.load()["framework"] == "trm"


# --- docx fill ---------------------------------------------------------------

def _seed(tmp_path, monkeypatch):
    engine = create_engine(f"sqlite:///{tmp_path/'t.db'}")
    SQLModel.metadata.create_all(engine)
    monkeypatch.setattr(db, "engine", engine)
    clauses = json.dumps([{"citation": "RMiT S 10.20", "clause_id": "10.20",
                           "section": "Patch Management", "binding": True, "score": 0.8,
                           "excerpt": "Apply patches promptly.", "reason": "Outdated component."}])
    with Session(engine) as s:
        a1 = Asset(ip_address="10.0.0.1", exposure=Exposure.external,
                   environment=Environment.production)
        a2 = Asset(ip_address="10.0.0.2", exposure=Exposure.internal,
                   environment=Environment.staging)
        s.add(a1); s.add(a2); s.commit(); s.refresh(a1); s.refresh(a2)
        # Same vuln name on two hosts -> one group, two instances.
        for a, port in ((a1, 443), (a2, 8443)):
            f = Finding(asset_id=a.id, tool="gvm", name="Outdated OpenSSL", port=port,
                        protocol="tcp", cve="CVE-2024-0001", solution="Upgrade OpenSSL.",
                        references="https://example.com/advisory", regulatory_clauses=clauses)
            s.add(f); s.commit(); s.refresh(f)
            s.add(FindingScore(finding_id=f.id, cvss_version="4.0", base_vector=_VEC,
                               base_score=9.8, base_severity="Critical", source="scan",
                               adj_vector=_VEC, adj_score=8.0, adj_severity="High",
                               fw_adj_score=9.1, fw_adj_severity="Critical"))
        # A second, distinct vuln (one host).
        f2 = Finding(asset_id=a1.id, tool="gvm", name="Weak TLS", port=443, protocol="tcp")
        s.add(f2); s.commit(); s.refresh(f2)
        s.add(FindingScore(finding_id=f2.id, cvss_version="4.0", base_vector=_VEC,
                           base_score=5.0, base_severity="Medium", source="scan",
                           adj_vector=_VEC, adj_score=5.0, adj_severity="Medium"))
        s.commit()


def _fill(tmp_path, monkeypatch, **kw):
    _seed(tmp_path, monkeypatch)
    # Skip the LibreOffice refresh/PDF step so the fill logic is tested deterministically
    # (the UNO refresh is exercised separately / manually).
    monkeypatch.setattr(docx_template, "_finalize", lambda docx_path, *, pdf: None)
    out = tmp_path / "report"
    meta = {"Client_Full_Name": "Acme Bank Berhad", "Client_Short_Name": "Acme",
            "Company_Full_Name": "SecTest Sdn Bhd", "Company_Short_Name": "SecTest"}
    written = docx_template.fill(TEMPLATE, out, framework="rmit", metadata=meta,
                                 target="10.0.0.0/24", provider="template", pdf=False, **kw)
    return written, Document(str(written[0]))


def _body_text(doc):
    return "".join((t.text or "") for t in doc.element.body.iter(qn("w:t")))


def test_fill_groups_by_vulnerability(tmp_path, monkeypatch):
    groups, counts, assets, fw = (None,) * 4
    _seed(tmp_path, monkeypatch)
    groups, counts, assets, fw = docx_template.load_model("rmit")
    names = [g.name for g in groups]
    assert names.count("Outdated OpenSSL") == 1           # 2 findings -> 1 group
    og = next(g for g in groups if g.name == "Outdated OpenSSL")
    assert len(og.instances) == 2                          # both hosts listed
    assert og.severity == "Critical" and og.cvss == "9.1"  # headline = worst fw_adj
    assert len(groups) == 2 and counts["Critical"] == 1 and counts["Medium"] == 1


def test_fill_resolves_docproperties_and_no_stray_tokens(tmp_path, monkeypatch):
    _, doc = _fill(tmp_path, monkeypatch)
    txt = _body_text(doc)
    assert "Acme Bank Berhad" in txt and "Client Full Name" not in txt
    assert "Bank Negara Malaysia RMiT" in txt              # Framework property set
    # Only the TOC field placeholder may remain (a consumer refreshes it).
    leftover = set(re.findall(r"\{\{[^}]+\}\}", txt))
    assert leftover <= {"{{F_TITLE}}"}


def test_fill_one_section_per_group_and_severity_shading(tmp_path, monkeypatch):
    written, doc = _fill(tmp_path, monkeypatch)
    import zipfile
    h2 = [p.text for p in doc.paragraphs if p.style and p.style.name == "Heading 2"]
    # 5 fixed H2 (scope/method/scoring/timeline/compliance) + 2 finding sections.
    assert "Outdated OpenSSL" in h2 and "Weak TLS" in h2
    xml = zipfile.ZipFile(str(written[0])).read("word/document.xml").decode()
    assert 'w:fill="C00000"' in xml and 'w:fill="FFC000"' in xml  # Critical + Medium shaded


def test_fill_sla_override_lands_in_table(tmp_path, monkeypatch):
    _, doc = _fill(tmp_path, monkeypatch, sla_overrides={"Critical": {"ext": 3, "int": 9}})
    txt = _body_text(doc)
    assert "within 3 days" in txt and "9 days" in txt


def test_resolve_template_missing_raises():
    import pytest
    with pytest.raises(docx_template.TemplateError):
        docx_template.resolve_template("no-such-template")


# --- report-quality fixes: env case, highlight, no-clause cell, steps restart ----

def test_title_env_capitalises():
    assert docx_template._title_env("production") == "Production"
    assert docx_template._title_env("uat") == "UAT"           # acronym
    assert docx_template._title_env("staging") == "Staging"


def _para_highlighted(doc, text_sub):
    for p in doc.element.body.iter(qn("w:p")):
        t = "".join((x.text or "") for x in p.iter(qn("w:t")))
        if text_sub in t:
            return any(r.find(qn("w:rPr")) is not None
                       and r.find(qn("w:rPr")).find(qn("w:highlight")) is not None
                       for r in p.findall(qn("w:r")))
    return None


def test_scope_env_is_capitalised_in_report(tmp_path, monkeypatch):
    _, doc = _fill(tmp_path, monkeypatch)
    txt = _body_text(doc)
    assert "Production" in txt and "Staging" in txt          # the two seeded assets


def test_filled_content_unhighlighted(tmp_path, monkeypatch):
    _, doc = _fill(tmp_path, monkeypatch)
    # auto-filled recommendation prose loses the authoring highlight …
    assert _para_highlighted(doc, "Patch outdated") is False
    # … and the former manual finding placeholders are now auto-filled (S5.5b) —
    # with no report inputs seeded they become the N/A note, also unhighlighted.
    assert _para_highlighted(doc, "non-intrusive assessment") is False


def test_no_clause_cell_is_clean(tmp_path, monkeypatch):
    _, doc = _fill(tmp_path, monkeypatch)          # 'Weak TLS' finding has no clauses
    txt = _body_text(doc)
    assert "No specific clause mapped" in txt
    assert "— - No specific" not in txt and " - No specific" not in txt


def test_steps_lists_restart_per_finding(tmp_path, monkeypatch):
    _, doc = _fill(tmp_path, monkeypatch)          # two findings -> two Steps lists
    ids = []
    for p in doc.element.body.iter(qn("w:p")):
        txt = "".join((x.text or "") for x in p.iter(qn("w:t")))
        ppr = p.find(qn("w:pPr"))
        npr = ppr.find(qn("w:numPr")) if ppr is not None else None
        # the filled {{F_STEPS}} list item (N/A note here — no steps seeded)
        if npr is not None and "non-intrusive assessment" in txt:
            nid = npr.find(qn("w:numId"))
            ids.append(nid.get(qn("w:val")) if nid is not None else None)
    assert len(ids) >= 2 and len(ids) == len(set(ids))    # own numId each -> restart at 1


def _png(path, w=16, h=16):
    """Write a genuinely valid RGB PNG (python-docx validates the header)."""
    import struct
    import zlib
    def ch(typ, data):
        c = typ + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xffffffff)
    raw = b"".join(b"\x00" + bytes((200, 60, 60)) * w for _ in range(h))
    path.write_bytes(b"\x89PNG\r\n\x1a\n" + ch(b"IHDR", struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
                     + ch(b"IDAT", zlib.compress(raw)) + ch(b"IEND", b""))


def test_manual_report_inputs_and_cover_filled(tmp_path, monkeypatch):
    """S5.5b: FindingReportInput (steps/comments/screenshot) + cover tokens fill."""
    import zipfile

    from finvap.models import FindingReportInput
    engine = create_engine(f"sqlite:///{tmp_path/'t.db'}")
    SQLModel.metadata.create_all(engine)
    monkeypatch.setattr(db, "engine", engine)
    shot = tmp_path / "poc.png"; _png(shot)
    with Session(engine) as s:
        a = Asset(ip_address="10.0.0.9", exposure=Exposure.external); s.add(a); s.commit(); s.refresh(a)
        f = Finding(asset_id=a.id, tool="gvm", name="Weak TLS", solution="Disable legacy TLS.")
        s.add(f); s.commit(); s.refresh(f)
        s.add(FindingScore(finding_id=f.id, cvss_version="4.0", base_vector=_VEC, base_score=7.0,
                           base_severity="High", source="scan", adj_vector=_VEC,
                           adj_score=7.0, adj_severity="High"))
        s.add(FindingReportInput(finding_id=f.id, poc_screenshot=str(shot),
                                 steps="Open a socket\nSend a TLS 1.0 hello\nHandshake succeeds",
                                 client_comments="Accepted, next patch cycle.",
                                 postverif_comments="Legacy TLS disabled."))
        s.commit()
    monkeypatch.setattr(docx_template, "_finalize", lambda docx_path, *, pdf: None)
    meta = {"Client_Full_Name": "Acme", "__draft_final": "Final",
            "__version": "1.0", "__cover_date": "04 Jul 2026"}
    written = docx_template.fill(TEMPLATE, tmp_path / "rep", framework="rmit",
                                 metadata=meta, provider="template", pdf=False)
    doc = Document(str(written[0])); body = _body_text(doc)
    assert "Open a socket" in body and "Handshake succeeds" in body   # steps lines
    assert "next patch cycle" in body                                  # client comments
    assert "Legacy TLS disabled" in body                               # post-verif comments
    assert "Final" in body and "1.0" in body and "04 Jul 2026" in body  # cover tokens
    with zipfile.ZipFile(written[0]) as z:
        assert any(n.startswith("word/media/") for n in z.namelist())  # screenshot embedded
    assert len(doc.inline_shapes) >= 1


def test_resolve_template_default_is_bundled():
    # No name + no config -> the bundled VA Template.docx.
    p = docx_template.resolve_template(None)
    assert p.name == "VA Template.docx" and p.exists()


# --- #1 grammar + #4 casing (pure helpers) -----------------------------------

def test_frag_makes_mid_sentence_fragment():
    assert docx_template._frag("The network has weak controls.") == "the network has weak controls"
    assert docx_template._frag("FTP cleartext logins.") == "FTP cleartext logins"  # keep acronym
    assert docx_template._frag("  Access sensitive data  ") == "access sensitive data"


def test_normalize_citation_casing():
    from finvap.compliance import normalize_citation
    assert normalize_citation("RMIT S 10.21, RMIT S 10.53") == "RMiT S 10.21, RMiT S 10.53"
    assert normalize_citation("TRM 7.4.1") == "TRM 7.4.1"          # unaffected


# --- #1 + #9: LLM prose (exec fragments + per-finding rewrite), no soffice ----

class _FakeLLM:
    name, model = "ollama", "fake"

    def available(self):
        return True, "fake"

    def complete(self, system, user, **kw):
        if "executive summary" in system:
            # deliberately echo the lead-in + capitalise + trailing period, to prove
            # the fragment cleanup (strip lead-in, lowercase, drop period).
            return ('{"posture": "requires moderate security improvement", '
                    '"core_issues": "The risks identified involve Outdated software components.", '
                    '"business_impacts": "allow an attacker to Access sensitive data.", '
                    '"recommendations": ["Patch the software.", "Harden the configuration."]}')
        return ('{"description": "A rewritten clear description.", '
                '"recommendation": "A rewritten clear recommendation."}')


def test_llm_prose_fragments_and_rewrite(tmp_path, monkeypatch):
    from finvap.reporting import prose
    _seed(tmp_path, monkeypatch)
    monkeypatch.setattr(docx_template, "_finalize", lambda docx_path, *, pdf: None)
    monkeypatch.setattr(docx_template, "get_provider", lambda *a, **k: _FakeLLM())
    monkeypatch.setattr(prose, "get_provider", lambda *a, **k: _FakeLLM())
    # Per-finding prose is now a persisted pipeline step, run before the report.
    assert prose.rewrite_findings(provider="ollama")["rewritten"] >= 1
    out = tmp_path / "report"
    written = docx_template.fill(TEMPLATE, out, framework="rmit",
                                 metadata={"Client_Full_Name": "Acme"}, target="10.0.0.0/24",
                                 provider="ollama", pdf=False)
    txt = _body_text(Document(str(written[0])))
    # exec fragments slot in grammatically: lead-in echo stripped, lowercased, no double period
    assert "involve outdated software components. If" in txt
    assert "risks identified involve the risks" not in txt      # echo removed
    assert "attacker to access sensitive data." in txt
    assert "attacker to allow an attacker" not in txt           # echo removed
    assert ".." not in txt.replace("...", "")                   # no stray double period
    # per-finding description/recommendation were rewritten (persisted) then rendered
    assert "A rewritten clear description." in txt
    assert "A rewritten clear recommendation." in txt


def test_rewrite_findings_persists_and_respects_override(tmp_path, monkeypatch):
    from sqlmodel import Session, select

    from finvap.models import Finding
    from finvap.reporting import prose
    _seed(tmp_path, monkeypatch)
    monkeypatch.setattr(prose, "get_provider", lambda *a, **k: _FakeLLM())
    # a human finalised "Weak TLS" — must be skipped by the AI prose step
    with Session(db.engine) as s:
        f = s.exec(select(Finding).where(Finding.name == "Weak TLS")).first()
        f.description = "HUMAN EDIT"; f.text_overridden = True
        s.add(f); s.commit()
    stats = prose.rewrite_findings(provider="ollama")
    assert stats["skipped_override"] >= 1
    with Session(db.engine) as s:
        outdated = s.exec(select(Finding).where(Finding.name == "Outdated OpenSSL")).first()
        weak = s.exec(select(Finding).where(Finding.name == "Weak TLS")).first()
        assert outdated.description == "A rewritten clear description."   # AI rewrote + saved
        assert weak.description == "HUMAN EDIT"                            # edit preserved
